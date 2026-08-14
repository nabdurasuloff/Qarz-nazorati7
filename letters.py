# -*- coding: utf-8 -*-
"""
Word shabloniga ma'lumotlarni joylab, tayyor xat (.docx) yaratish.
"""
import os
import re
import sys
import datetime
from docx import Document


def _base_dir():
    # PyInstaller --onefile bilan yig'ilganda fayllar vaqtinchalik papkaga
    # (sys._MEIPASS) ochiladi; oddiy ishga tushirishda esa shu faylning papkasi.
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))


TEMPLATE_PATH = os.path.join(_base_dir(), 'templates', 'xat_shablon.docx')
DAVO_ARIZA_TEMPLATE_PATH = os.path.join(_base_dir(), 'templates', 'davo_ariza_shablon.docx')


def _fmt_summa(val):
    try:
        val = float(val)
    except (TypeError, ValueError):
        return str(val)
    return f"{val:,.0f}".replace(',', ' ')


def _replace_in_paragraph(paragraph, mapping):
    full_text = ''.join(run.text for run in paragraph.runs)
    if '{{' not in full_text:
        return
    new_text = full_text
    for key, val in mapping.items():
        new_text = new_text.replace('{{' + key + '}}', str(val))
    if new_text == full_text:
        return
    # Barcha runlarni tozalab, birinchi runga yangi matnni yozamiz
    # (formatlashni birinchi run’dan olamiz)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)


def _replace_everywhere(doc, mapping):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)


def generate_letter(output_path, xat_turi, mijoz_ism, mijoz_manzil, portfel_row, settings,
                     anketa_raqami=None, rahbar_ism=None):
    """
    xat_turi: 'Ogohlantirish' yoki 'Talabnoma'
    portfel_row: dict — bitta portfel qatoridagi ma'lumot (database.py formatida)
    settings: dict — get_all_settings() natijasi
    """
    doc = Document(TEMPLATE_PATH)

    sarlavha = 'ОГОҲЛАНТИРИШ ХАТИ' if xat_turi == 'Ogohlantirish' else 'ТАЛАБНОМА'

    holat_sanasi = datetime.date.today().strftime('%d.%m.%Y')

    mapping = {
        'BANK_NOMI': settings.get('bank_nomi', ''),
        'BANK_QISQA_NOMI': settings.get('bank_qisqa_nomi', ''),
        'BANK_MANZIL': settings.get('bank_manzil', ''),
        'BANK_EMAIL': settings.get('bank_email', ''),
        'BANK_SAYT': settings.get('bank_sayt', ''),
        'BANK_TEL': settings.get('bank_tel', ''),
        'BANK_MOBIL_ILOVA': settings.get('bank_mobil_ilova', ''),
        'BANK_KODI': settings.get('bank_kodi', ''),
        'ALOQA_MARKAZI_TEL': settings.get('aloqa_markazi_tel', ''),
        'FILIAL_NOMI': settings.get('filial_nomi', ''),
        'FILIAL_TEL': settings.get('filial_tel', ''),
        'RAHBAR_ISM': rahbar_ism or settings.get('rahbar_ism', ''),

        'MIJOZ_ISM': mijoz_ism or '',
        'MIJOZ_MANZIL': mijoz_manzil or '',
        'SARLAVHA': sarlavha,

        'KREDIT_SUMMA': _fmt_summa(portfel_row.get('ead', 0)),
        'KREDIT_MAQSAD': portfel_row.get('tulov_maqsadi', '') or '',
        'HOLAT_SANASI': holat_sanasi,
        'JAMI_QARZ': _fmt_summa(portfel_row.get('jami_qarz', 0)),
        'ASOSIY_QARZ': _fmt_summa(portfel_row.get('asosiy_qarz', 0)),
        'FOIZ_QARZ': _fmt_summa(portfel_row.get('foiz_qarz', 0)),
        'JARIMA': _fmt_summa(portfel_row.get('jarima', 0)),
        'TOLOV_MUDDATI': f"{settings.get('tolov_muddati_kun', '10')} банк иш куни",
        'ANKETA_RAQAM': anketa_raqami or portfel_row.get('anketa_raqami', ''),
    }

    _replace_everywhere(doc, mapping)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_davo_ariza(*args, **kwargs):
    raise NotImplementedError(
        "Bu funksiya eskirgan. O'rniga generate_davo_ariza_v2() dan foydalaning."
    )


def safe_filename(text):
    text = re.sub(r'[\\/*?:"<>|]', '', str(text))
    text = text.strip().replace(' ', '_')
    return text[:80]


DAVO_ARIZA_TEMPLATES = {
    'jismoniy_oddiy': 'davo_jismoniy_oddiy.docx',
    'jismoniy_kafil': 'davo_jismoniy_kafil.docx',
    'yuridik_oddiy': 'davo_yuridik_oddiy.docx',
    'yuridik_kafil': 'davo_yuridik_kafil.docx',
    'yuridik_kafil_garov': 'davo_yuridik_kafil_garov.docx',
    'muddatidan_oldin': 'davo_muddatidan_oldin.docx',
}

DAVO_ARIZA_NOMLARI = {
    'jismoniy_oddiy': "Jismoniy shaxs — oddiy (kafilsiz)",
    'jismoniy_kafil': "Jismoniy shaxs — kafil bilan",
    'yuridik_oddiy': "Yuridik/F.X — oddiy (kafilsiz)",
    'yuridik_kafil': "Yuridik shaxs — kafil bilan (foiz-penya)",
    'yuridik_kafil_garov': "Yuridik shaxs — kafil + garov mulkiga qaratish",
    'muddatidan_oldin': "Shartnomani bekor qilish (muddatidan oldin undirish)",
}


def _oy_farqi(sana1_str, sana2_str):
    """Ikki sana orasidagi farqni oy hisobida qaytaradi (taxminiy)."""
    for fmt_pair in [('%d.%m.%Y', '%d.%m.%Y')]:
        try:
            d1 = datetime.datetime.strptime(str(sana1_str)[:10], '%d.%m.%Y')
            d2 = datetime.datetime.strptime(str(sana2_str)[:10], '%d.%m.%Y')
            months = (d2.year - d1.year) * 12 + (d2.month - d1.month)
            return max(months, 0)
        except (ValueError, TypeError):
            return ''
    return ''


def generate_davo_ariza_v2(turi, output_path, portfel_row, mijoz, taminot, settings,
                            xat_sanasi='', xat_turi_nomi='', imzo_ism_override=None):
    """
    turi: DAVO_ARIZA_TEMPLATES kalitlaridan biri.
    mijoz: mijozlar jadvalidagi dict (yoki None).
    taminot: davo_taminot jadvalidagi dict (yoki None) — kafil/garov ma'lumotlari.
    """
    tpl_file = DAVO_ARIZA_TEMPLATES.get(turi)
    if not tpl_file:
        raise ValueError(f"Noma'lum davo ariza turi: {turi}")
    tpl_path = os.path.join(_base_dir(), 'templates', tpl_file)
    doc = Document(tpl_path)

    mijoz = mijoz or {}
    taminot = taminot or {}
    holat_sanasi = datetime.date.today().strftime('%d.%m.%Y')
    muddat_oy = _oy_farqi(portfel_row.get('shartnoma_sanasi'), portfel_row.get('shartnoma_tugash_sanasi'))

    jami_ead = float(portfel_row.get('ead', 0) or 0)
    asosiy = float(portfel_row.get('asosiy_qarz', 0) or 0)
    foiz = float(portfel_row.get('foiz_qarz', 0) or 0)
    jarima_v = float(portfel_row.get('jarima', 0) or 0)
    jami = asosiy + foiz + jarima_v
    muddati_kelmagan = max(jami_ead - jami, 0)

    pochta = taminot.get('pochta_xarajati') or settings.get('pochta_xarajati_standart', '41200')

    mapping = {
        'ARIZA_SANA_QATORI': f'{holat_sanasi}-yil',
        'SUD_NOMI': settings.get('sud_iqtisodiy_nomi' if turi.startswith('yuridik') or turi == 'muddatidan_oldin'
                                  else 'sud_fuqarolik_nomi', ''),
        'PALATA_NOMI': settings.get('palata_nomi', ''),

        'BANK_NOMI': settings.get('bank_nomi', ''),
        'BANK_QISQA_NOMI': settings.get('bank_qisqa_nomi', ''),
        'BANK_STIR': settings.get('bank_stir', ''),
        'BANK_KODI_BOSH': settings.get('bank_kodi_bosh', ''),
        'BANK_HISOB_RAQAM_BOSH': settings.get('bank_hisob_raqami_bosh', ''),
        'BANK_MANZIL_BOSH': settings.get('bank_manzil', ''),
        'BANK_MANZIL_FILIAL': settings.get('bank_rasmiy_manzil_filial', ''),
        'BANK_REKVIZIT_QATORI': (f"ҳ|р: {settings.get('bank_hisob_raqami_filial', '')}, "
                                  f"банк коди: {settings.get('bank_kodi_filial', '')}, "
                                  f"СТИР: {settings.get('bank_stir', '')}, "
                                  f"Манзил: {settings.get('bank_rasmiy_manzil_filial', '')}"),
        'FILIAL_NOMI': settings.get('filial_nomi', ''),

        'MIJOZ_ISM': mijoz.get('ism') or portfel_row.get('mijoz_nomi', ''),
        'MIJOZ_MANZIL': mijoz.get('manzil', '') or '',
        'MIJOZ_PINFL': portfel_row.get('pinfl', '') or mijoz.get('hujjat_raqami', '') or '',
        'MIJOZ_STIR': portfel_row.get('stir', '') or '',
        'MIJOZ_BANK_KODI': portfel_row.get('filial_kodi', '') or '',
        'MIJOZ_HISOB_RAQAM': portfel_row.get('kredit_hisob_raqami', '') or '',
        'MIJOZ_RAHBAR': mijoz.get('rahbar_ism', '') or '',
        'MIJOZ_TEL': mijoz.get('telefon', '') or '',
        'MIJOZ_PASSPORT_TOLIQ': _passport_toliq(mijoz),
        'BANK_HUDUDIY_NOMI': f"{settings.get('viloyat_nomi','Сирдарё')} вилоят худудий бошкармаси ({settings.get('filial_nomi','')} филиали)",

        'KAFIL_ISM': taminot.get('kafil_ism', '') or '',
        'KAFIL_MANZIL': taminot.get('kafil_manzil', '') or '',
        'KAFIL_PINFL': taminot.get('kafil_pinfl', '') or '',
        'KAFIL_TEL': taminot.get('kafil_tel', '') or '',
        'KAFIL_TUGILGAN_SANA': taminot.get('kafil_passport_sana', '') or '',
        'KAFIL_TUGILGAN_JOY': taminot.get('kafil_passport_organ', '') or '',
        'KAFIL_PASSPORT_TOLIQ': (
            f"{taminot.get('kafil_passport','')}, {taminot.get('kafil_passport_sana','')} йил "
            f"{taminot.get('kafil_passport_organ','')}дан берилган" if taminot.get('kafil_passport') else ''
        ),

        'GAROV_TAVSIFI': taminot.get('garov_tavsifi', '') or '',
        'GAROV_BAHOSI': _fmt_summa(taminot.get('garov_bahosi', 0)),

        'SHARTNOMA_SANA': portfel_row.get('shartnoma_sanasi', '') or '',
        'KREDIT_MUDDATI_OY': muddat_oy,
        'IMTIYOZLI_DAVR_OY': '',
        'YILLIK_FOIZ': portfel_row.get('yillik_foiz', '') or '',
        'KREDIT_MAQSAD': portfel_row.get('tulov_maqsadi', '') or '',
        'KREDIT_SUMMA': _fmt_summa(portfel_row.get('ead', 0)),

        'HOLAT_SANASI': holat_sanasi,
        'JAMI_QARZ': _fmt_summa(jami),
        'ASOSIY_QARZ': _fmt_summa(asosiy),
        'FOIZ_QARZ': _fmt_summa(foiz),
        'JARIMA': _fmt_summa(jarima_v),
        'MUDDATI_KELMAGAN_ASOSIY': _fmt_summa(muddati_kelmagan),
        'POCHTA_XARAJATI': _fmt_summa(pochta),

        'XAT_SANASI': xat_sanasi,
        'XAT_TURI_NOMI': xat_turi_nomi,
        'VOQEALAR_TAVSIFI': taminot.get('garov_tavsifi', '') or '[Voqealar tavsifini shu yerga kiriting]',

        'IMZO_ISM': imzo_ism_override or settings.get('sud_ariza_imzo_ism', '') or '',
        'IMZO_LAVOZIM': settings.get('sud_ariza_imzo_lavozimi', ''),
    }

    _replace_everywhere(doc, mapping)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _passport_toliq(mijoz):
    if not mijoz:
        return ''
    hujjat = mijoz.get('hujjat_raqami', '') or ''
    sana = mijoz.get('passport_sana', '') or ''
    organ = mijoz.get('passport_organ', '') or ''
    if not hujjat:
        return ''
    parts = [hujjat]
    if sana:
        parts.append(f"{sana} йил")
    if organ:
        parts.append(f"{organ}дан берилган")
    return ', '.join(parts) if len(parts) > 1 else parts[0]


def convert_docx_to_pdf(docx_path, pdf_path=None, delete_docx=False):
    """
    .docx faylni .pdf ga aylantiradi. Windows'da MS Word o'rnatilgan bo'lishi
    shart (docx2pdf shu orqali ishlaydi). Agar Word bo'lmasa, xato chiqadi.
    """
    if pdf_path is None:
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
    try:
        from docx2pdf import convert
    except ImportError:
        raise RuntimeError(
            "PDF yaratish uchun 'docx2pdf' kutubxonasi o'rnatilmagan. "
            "requirements.txt orqali o'rnating: pip install docx2pdf pywin32"
        )
    try:
        convert(docx_path, pdf_path)
    except Exception as e:
        raise RuntimeError(
            f"PDF'ga aylantirishda xato: {e}\n"
            "Bu funksiya faqat Windows'da, Microsoft Word o'rnatilgan bo'lsa ishlaydi."
        )
    if delete_docx and os.path.exists(docx_path):
        os.remove(docx_path)
    return pdf_path
